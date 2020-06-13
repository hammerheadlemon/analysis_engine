'''
Programme that compiles project dashboards/summary sheets.

Input:
1) four quarters worth of data

Output:
1) MS word document in structure of summary sheet / dashboard - with some areas missing, see below.

Supplementary programmes that need to be run to build charts for summary pages. Charts should be built and cut and paste
into dashboards/summary sheets:
1) project_financial_profile.py . For financial charts
2) milestone_comparison_3_qrts_proj.py . For milestones tables

'''

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, RGBColor, Inches, Pt
from docx.enum.section import WD_SECTION_START, WD_ORIENT
import difflib
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib import rcParams
import numpy as np
import datetime
from datetime import timedelta
from textwrap import wrap


from analysis.engine_functions import convert_rag_text, all_milestone_data_bulk, project_time_difference, \
    all_milestones_dict
from analysis.data import list_of_masters_all, root_path, latest_cost_profiles, last_cost_profiles, \
    baseline_1_cost_profiles, year_list, SRO_conf_key_list, SRO_conf_table_list, bc_index, p_current_milestones, \
    p_last_milestones, first_diff_data, ipdc_date, \
    a66, a303, crossrail, thameslink, south_west_route_capacity

import os

milestone_filter_start_date = ipdc_date - timedelta(days=30*6)
milestone_filter_end_date = ipdc_date + timedelta(days=545)

def cell_colouring(cell, colour):
    '''
    Function that handles cell colouring
    cell: cell reference
    color: colour reference
    '''

    try:
        if colour == 'R':
            colour = parse_xml(r'<w:shd {} w:fill="cb1f00"/>'.format(nsdecls('w')))
        elif colour == 'A/R':
            colour = parse_xml(r'<w:shd {} w:fill="f97b31"/>'.format(nsdecls('w')))
        elif colour == 'A':
            colour = parse_xml(r'<w:shd {} w:fill="fce553"/>'.format(nsdecls('w')))
        elif colour == 'A/G':
            colour = parse_xml(r'<w:shd {} w:fill="a5b700"/>'.format(nsdecls('w')))
        elif colour == 'G':
            colour = parse_xml(r'<w:shd {} w:fill="17960c"/>'.format(nsdecls('w')))

        cell._tc.get_or_add_tcPr().append(colour)

    except TypeError:
        pass

def compare_text_showall(text_1, text_2, doc):
    '''
    Function places text into doc highlighting all changes.
    text_1: latest text. string.
    text_2: last text. string
    doc: word doc
    '''

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.strike = True
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i])

    return doc

def compare_text_newandold(text_1, text_2, doc):
    '''
    function that places text into doc highlighting new and old text
    text_1: latest text. string.
    text_2: last text. string
    doc: word doc
    '''

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            pass
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i][1:])

    return doc

def produce_word_doc():
    '''Function that compiles each summary sheet'''

    master_list = list_of_masters_all[0:4]
    test_project_list = [crossrail]
                         #south_west_route_capacity, a66, a303, crossrail, thameslink]

    for project_name in test_project_list:
        doc = Document()
        print(project_name)

        font = doc.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(10)

        heading = str(project_name)
        intro = doc.add_heading(str(heading), 0)
        intro.alignment = 1
        intro.bold = True

        para_1 = doc.add_paragraph()
        sro_name = list_of_masters_all[0].data[project_name]['Senior Responsible Owner (SRO)']
        if sro_name is None:
            sro_name = 'TBC'

        sro_phone = list_of_masters_all[0].data[project_name]['SRO Phone No.']
        if sro_phone == None:
            sro_phone = 'TBC'

        para_1.add_run('SRO name:  ' + str(sro_name) + ',   Tel:  ' + str(sro_phone))

        para_2 = doc.add_paragraph()
        pd_name = list_of_masters_all[0].data[project_name]['Project Director (PD)']
        if pd_name is None:
            pd_name = 'TBC'

        pd_phone = list_of_masters_all[0].data[project_name]['PD Phone No.']
        if pd_phone is None:
            pd_phone = 'TBC'

        para_2.add_run('PD name:  ' + str(pd_name) + ',   Tel:  ' + str(pd_phone))

        '''Start of table with DCA confidence ratings'''
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Delivery confidence'
        hdr_cells[1].text = 'This quarter'
        hdr_cells[2].text = str(list_of_masters_all[1].quarter)
        hdr_cells[3].text = str(list_of_masters_all[2].quarter)
        hdr_cells[4].text = str(list_of_masters_all[3].quarter)

        #RAG ratings
        for x, dca_key in enumerate(SRO_conf_key_list):
            row_cells = table.add_row().cells
            row_cells[0].text = dca_key
            for i, master in enumerate(master_list):
                try:
                    rating = convert_rag_text(master.data[project_name][dca_key])
                    row_cells[i+1].text = rating
                    cell_colouring(row_cells[i+1], rating)
                except (KeyError, TypeError):
                    row_cells[i].text = 'N/A'

        table.style = 'Table Grid'

        # TODO develop way of setting table column widths
        # TODO add text re what red text means.
        #set_col_widths(table)

        #DCA narrative
        doc.add_paragraph()

        headings_list = ['SRO delivery confidence narrative',
                         'Financial cost narrative',
                         'Financial comparison with last quarter',
                         'Financial comparison this quarter',
                         'Benefits Narrative',
                         'Benefits comparison with last quarter',
                         'Benefits comparison this quarter',
                         'Milestone narrative']

        narrative_keys_list = ['Departmental DCA Narrative',
                               'Project Costs Narrative',
                               'Cost comparison with last quarters cost narrative',
                               'Cost comparison within this quarters cost narrative',
                               'Benefits Narrative',
                               'Ben comparison with last quarters cost - narrative',
                               'Ben comparison within this quarters cost - narrative',
                               'Milestone Commentary']

        for x in range(len(headings_list)):
            doc.add_paragraph().add_run(str(headings_list[x])).bold = True
            text_one = str(list_of_masters_all[0].data[project_name][narrative_keys_list[x]])
            try:
                text_two = str(list_of_masters_all[1].data[project_name][narrative_keys_list[x]])
            except KeyError:
                text_two = text_one

            #different options for comparing costs
            # compare_text_showall(dca_a, dca_b, doc)
            compare_text_newandold(text_one, text_two, doc)

        '''Financial data'''
        fin_all = get_financial_totals(project_name) # all totals

        #totals by spent, profiled and unprofiled
        total_fin = fin_all[0]
        total_profiled_bl = total_fin[1] - (total_fin[0] + total_fin[2])
        total_profiled_lst = total_fin[4] - (total_fin[3] + total_fin[5])
        total_profiled_now = total_fin[7] - (total_fin[6] + total_fin[8])
        t_spent = np.array([total_fin[0], total_fin[3], total_fin[6]])
        t_profiled = np.array([total_profiled_bl, total_profiled_lst, total_profiled_now])
        t_unprofiled = np.array([total_fin[2], total_fin[5], total_fin[8]])

        #totals for rdel and cdel
        rdel_fin = fin_all[1]
        cdel_fin = fin_all[2]
        rdel_profiled = rdel_fin[7] - (rdel_fin[6] + rdel_fin[8])
        cdel_profiled = cdel_fin[7] - (cdel_fin[6] + cdel_fin[8])
        rc_spent = np.array([rdel_fin[6], cdel_fin[6]])
        rc_profiled = np.array([rdel_profiled, cdel_profiled])
        rc_unprofiled = np.array([rdel_fin[8], cdel_fin[8]])

        # '''Financial Meta data'''
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'WLC'
        hdr_cells[1].text = 'Spent'
        hdr_cells[2].text = 'Profiled'
        hdr_cells[3].text = 'Unprofiled'
        hdr_cells[4].text = 'RDEl Total'
        hdr_cells[5].text = 'CDEl Total'
        row_cells = table.add_row().cells
        row_cells[0].text = str(total_fin[7])
        row_cells[1].text = str(total_fin[6])
        row_cells[2].text = str(total_profiled_now)
        row_cells[3].text = str(total_fin[8])
        row_cells[4].text = str(rdel_fin[7])
        row_cells[5].text = str(cdel_fin[7])

        # wlc = round(list_of_masters_all[0].data[project_name]['Total Forecast'], 1)
        # table1.cell(1, 0).text = str(wlc)
        # # str(list_of_masters_all[0].data[project_name]['Total Forecast'])
        # #a = list_of_masters_all[0].data[project_name]['Total Forecast']
        # b = list_of_masters_all[0].data[project_name]['Pre 19-20 RDEL Forecast Total']
        # if b == None:
        #     b = 0
        # c = list_of_masters_all[0].data[project_name]['Pre 19-20 CDEL Forecast Total']
        # if c == None:
        #     c = 0
        # d = list_of_masters_all[0].data[project_name]['Pre 19-20 Forecast Non-Gov']
        # if d == None:
        #     d = 0
        # e = b + c + d
        # try:
        #     c = round(e / wlc * 100, 1)
        # except (ZeroDivisionError, TypeError):
        #     c = 0
        # table1.cell(1, 1).text = str(c) + '%'
        # a = str(list_of_masters_all[0].data[project_name]['Source of Finance'])
        # b = list_of_masters_all[0].data[project_name]['Other Finance type Description']
        # if b == None:
        #     table1.cell(1, 2).text = a
        # else:
        #     table1.cell(1, 2).text = a + ' ' + str(b)
        # table1.cell(1, 3).text = str(list_of_masters_all[0].data[project_name]['Real or Nominal - Actual/Forecast'])
        # table1.cell(1, 4).text = ''

        '''start of analysis'''
        new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        new_width, new_height = new_section.page_height, new_section.page_width
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height

        '''Financial charts'''
        #doc.add_paragraph().add_run(str('Financial Analysis')).bold = True #don't need this
        financial_graph_charts(doc, project_name, t_spent, t_profiled, t_unprofiled, total_fin,
                               rc_spent, rc_profiled, rc_unprofiled)

        '''milestone swimlane charts'''
        #chart of with milestone over the next two years
        p_baseline_milestones = all_milestones_dict([project_name], list_of_masters_all[bc_index[project_name][2]])
        m_data = milestone_schedule_data_filtered(p_current_milestones, p_last_milestones, p_baseline_milestones,
                                         project_name)
        # add \n to y axis labels and cut down if two long
        labels = ['\n'.join(wrap(l, 40)) for l in m_data[0]]
        final_labels = []
        for l in labels:
            if len(l) > 70:
                final_labels.append(l[:70])
            else:
                final_labels.append(l)

        #Chart
        no_milestones = len(m_data[0])

        if no_milestones <= 15:
            title = project_name + ' schedule next two years'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels), np.array(m_data[1]), np.array(m_data[2]), \
                                  np.array(m_data[3]), title)

        if 16 <= no_milestones <= 30:
            half = int(no_milestones/2)
            title = project_name + ' schedule next two years'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[:half]), np.array(m_data[1][:half]),
                                      np.array(m_data[2][:half]), np.array(m_data[3][:half]), title)
            title = project_name + ' schedule next two years Cont'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[half:no_milestones]),
                                      np.array(m_data[1][half:no_milestones]),
                                      np.array(m_data[2][half:no_milestones]),
                                      np.array(m_data[3][half:no_milestones]), title)

        #chart with all milestones
        p_baseline_milestones = all_milestones_dict([project_name], list_of_masters_all[bc_index[project_name][2]])
        m_data = milestone_schedule_data(p_current_milestones, p_last_milestones, p_baseline_milestones,
                                                  project_name)
        #print(m_data)
        # add \n to y axis labels and cut down if two long
        labels = ['\n'.join(wrap(l, 40)) for l in m_data[0]]
        final_labels = []
        for l in labels:
            if len(l) > 70:
                final_labels.append(l[:70])
            else:
                final_labels.append(l)

        # Chart
        no_milestones = len(m_data[0])

        if no_milestones <= 20:
            title = project_name + ' schedule all'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels), np.array(m_data[1]),
                                      np.array(m_data[2]), \
                                      np.array(m_data[3]), title)

        if 21 <= no_milestones <= 40:
            half = int(no_milestones/2)
            title = project_name + ' schedule all'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[:half]),
                                      np.array(m_data[1][:half]),
                                      np.array(m_data[2][:half]), np.array(m_data[3][:half]), title)
            title = project_name + ' schedule all Cont'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[half:no_milestones]),
                                      np.array(m_data[1][half:no_milestones]),
                                      np.array(m_data[2][half:no_milestones]),
                                      np.array(m_data[3][half:no_milestones]), title)

        if 41 <= no_milestones <= 60:
            third = int(no_milestones/3)
            title = project_name + ' schedule all'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[:third]),
                                      np.array(m_data[1][:third]),
                                      np.array(m_data[2][:third]), np.array(m_data[3][:third]), title)
            title = project_name + ' schedule all Cont'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[third:third*2]),
                                      np.array(m_data[1][third:third*2]),
                                      np.array(m_data[2][third:third*2]),
                                      np.array(m_data[3][third:third*2]), title)
            title = project_name + ' schedule all Cont'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[third*2:no_milestones]),
                                      np.array(m_data[1][third*2:no_milestones]),
                                      np.array(m_data[2][third*2:no_milestones]),
                                      np.array(m_data[3][third*2:no_milestones]), title)


        #milestone table
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        #table heading
        doc.add_paragraph().add_run(str('Project high-level milestones')).bold = True
        #doc.add_paragraph()
        some_text = 'The below table presents all project reported remaining high-level milestones, with six months grace ' \
                    'from close of the current quarter. Milestones are sorted in chronological order. Changes in milestones' \
                    ' dates in comparison to last quarter and baseline have been calculated and are provided.'
        doc.add_paragraph().add_run(str(some_text)).italic = True

        milestone_table(doc, p_baseline_milestones, project_name)



        # '''milestone section'''
        # y = doc.add_paragraph()
        # heading = 'Planning information'
        # y.add_run(str(heading)).bold = True
        #
        # '''Milestone Meta data'''
        # table1 = doc.add_table(rows=2, cols=4)
        # table1.cell(0, 0).text = 'Project Start Date:'
        # table1.cell(0, 1).text = 'Latest Approved Business Case:'
        # table1.cell(0, 2).text = 'Start of Operations:'
        # table1.cell(0, 3).text = 'Project End Date:'
        #
        # key_dates = milestone_master[project_name]
        #
        # #c = key_dates['Start of Project']
        # try:
        #     c = tuple(key_dates['Start of Project'])[0]
        #     c = datetime.datetime.strptime(c.isoformat(), '%Y-%M-%d').strftime('%d/%M/%Y')
        # except (KeyError, AttributeError):
        #     c = 'Not reported'
        #
        # table1.cell(1, 0).text = str(c)
        #
        # table1.cell(1, 1).text = str(list_of_masters_all[0].data[project_name]['IPDC approval point'])
        #
        # try:
        #     a = tuple(key_dates['Start of Operation'])[0]
        #     a = datetime.datetime.strptime(a.isoformat(), '%Y-%M-%d').strftime('%d/%M/%Y')
        #     table1.cell(1, 2).text = str(a)
        # except (KeyError, AttributeError):
        #     table1.cell(1, 2).text = 'Not reported'
        #
        # #b = key_dates['Project End Date']
        # try:
        #     b = tuple(key_dates['Project End Date'])[0]
        #     b = datetime.datetime.strptime(b.isoformat(), '%Y-%M-%d').strftime('%d/%M/%Y')
        # except (KeyError, AttributeError):
        #     b = 'Not reported'
        # table1.cell(1, 3).text = str(b)
        #
        # # TODO: workout generally styling options for doc, paragraphs and tables
        #
        # '''milestone narrative text'''
        # doc.add_paragraph()
        # y = doc.add_paragraph()
        # heading = 'SRO Milestone Narrative'
        # y.add_run(str(heading)).bold = True
        #
        # mile_dca_a = list_of_masters_all[0].data[project_name]['Milestone Commentary']
        # if mile_dca_a == None:
        #     mile_dca_a = 'None'
        #
        # try:
        #     mile_dca_b = list_of_masters_all[1].data[project_name]['Milestone Commentary']
        #     if mile_dca_b == None:
        #         mile_dca_b = 'None'
        # except KeyError:
        #     mile_dca_b = mile_dca_a
        #
        # # compare_text_showall()
        # compare_text_newandold(mile_dca_a, mile_dca_b, doc)
        #
        # '''milestone chart heading'''
        # y = doc.add_paragraph()
        # heading = 'Project reported high-level milestones and schedule changes'
        # y.add_run(str(heading)).bold = True
        # y = doc.add_paragraph()
        # some_text = 'The below table presents all project reported remaining high-level milestones, with six months grace ' \
        #             'from close of the current quarter. Milestones are sorted in chronological order. Changes in milestones' \
        #             ' dates in comparison to last quarter and baseline have been calculated and are provided.'
        # y.add_run(str(some_text)).italic = True
        # y = doc.add_paragraph()
        # y.add_run('{insert chart}')

        #TODO add quarter info in title
        doc.save(root_path/'output/{}_summary.docx'.format(project_name))

def no_graphs(num_milestones, num):
    if num_milestones <= num:
        return 1
    if num+1 <= num_milestones <= num*2:
        return 2
    if (num*2)+1 <= num_milestones <= num*3:  #side note >= 22.5 returns false for some reason.
        return 3
    if (num*3)+1 <= num_milestones:
        return 4

def combine_narrtives(project_name, master, key_list):
    '''function that combines text across different keys'''
    output = ''
    for key in key_list:
        output = output + str(master.data[project_name][key])

    return output

amended_year_list = year_list[:-1]

def set_col_widths(table):
    widths = (Inches(2), Inches(1), Inches(1), Inches(1), Inches(1))
    for col in table.columns:
        #print(col)
        for idx, width in enumerate(widths):
            col.cells[idx].width = width

def get_financial_profile(project_name, cost_type):
    '''gets project financial data'''
    latest = []
    last = []
    baseline = []
    for year in amended_year_list:
        baseline.append(baseline_1_cost_profiles[project_name][year + cost_type])
        last.append(last_cost_profiles[project_name][year + cost_type])
        latest.append(latest_cost_profiles[project_name][year + cost_type])

    return latest, last, baseline

def financial_graph_charts(doc, project_name, total_spent, total_profiled, total_unprofiled, total_fin,
                           rc_spent, rc_profiled, rc_unprofiled):

    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2) #four sub plotsprint

    '''cost profile charts'''
    profile_data_total = get_financial_profile(project_name, ' total')
    profile_data_rdel = get_financial_profile(project_name, ' RDEL Forecast Total')
    profile_data_cdel = get_financial_profile(project_name, ' CDEL Forecast Total')

    year = ['19/20', '20/21', '21/22', '22/23', '23/24', '24/25', '25/26', '26/27', '27/28', '28/29']
    baseline_profile_total = profile_data_total[2]
    last_profile_total = profile_data_total[1]
    latest_profile_total = profile_data_total[0]

    latest_profile_rdel = profile_data_rdel[0]
    latest_profile_cdel = profile_data_cdel[0]

    fig.suptitle(project_name + ' financial analysis', fontweight='bold') # title
    # set fig size
    #fig.set_figheight(6)
    #fig.set_figwidth(8)

    #plot cost change profile chart
    ax1.plot(year, baseline_profile_total, label='Baseline', linewidth=3.0, marker="o")
    ax1.plot(year, last_profile_total, label='Last quarter', linewidth=3.0, marker="o")
    ax1.plot(year, latest_profile_total, label='Latest', linewidth=3.0, marker="o")

    #cost profile change chart styling
    ax1.tick_params(axis='x', which='major', labelsize=6)
    ax1.set_ylabel('Cost (£m)')
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style('italic')
    ylab1.set_size(8)
    ax1.grid(color='grey', linestyle='-', linewidth=0.2)
    ax1.legend()
    ax1.set_title('Fig 1 - cost profile changes', loc='left', fontsize=8, fontweight='bold')

    # plot rdel/cdel chart data
    ax3.plot(year, latest_profile_cdel, label='CDEL', linewidth=3.0, marker="o")
    ax3.plot(year, latest_profile_rdel, label='RDEL', linewidth=3.0, marker="o")

    #rdel/cdel profile chart styling
    ax3.tick_params(axis='x', which='major', labelsize=6)
    ax3.set_xlabel('Financial Years')
    ax3.set_ylabel('Cost (£m)')
    xlab3 = ax3.xaxis.get_label()
    ylab3 = ax3.yaxis.get_label()
    xlab3.set_style('italic')
    xlab3.set_size(8)
    ylab3.set_style('italic')
    ylab3.set_size(8)
    ax3.grid(color='grey', linestyle='-', linewidth=0.2)
    ax3.legend()

    ax3.set_title('Fig 2 - cost profile spend type', loc='left', fontsize=8, fontweight='bold')

    #Spent, Profiled and Unprofile chart
    labels = ['Baseline', 'Last Quarter', 'Latest']
    width = 0.5
    ax2.bar(labels, total_spent, width, label='Spent')
    ax2.bar(labels, total_profiled, width, bottom=total_spent, label='Profiled')
    ax2.bar(labels, total_unprofiled, width, bottom=total_spent + total_profiled, label='Unprofiled')
    ax2.legend()
    ax2.set_ylabel('Cost (£m)')
    ylab2 = ax2.yaxis.get_label()
    ylab2.set_style('italic')
    ylab2.set_size(8)
    ax2.tick_params(axis='x', which='major', labelsize=6)
    ax2.tick_params(axis='y', which='major', labelsize=6)
    ax2.set_title('Fig 3 - wlc break down', loc='left', fontsize=8, fontweight='bold')

    #scaling y axis
    y_max = max(total_fin) + max(total_fin)*1/5
    ax2.set_ylim(0, y_max)

    #rdel/cdel bar chart

    labels = ['RDEL', 'CDEL']
    width = 0.5
    ax4.bar(labels, rc_spent, width, label='Spent')
    ax4.bar(labels, rc_profiled, width, bottom=rc_spent, label='Profiled')
    ax4.bar(labels, rc_unprofiled, width, bottom=rc_spent + rc_profiled, label='Unprofiled')
    ax4.legend()
    ax4.set_ylabel('Costs (£m)')
    ylab3 = ax4.yaxis.get_label()
    ylab3.set_style('italic')
    ylab3.set_size(8)
    ax4.tick_params(axis='x', which='major', labelsize=6)
    ax4.tick_params(axis='y', which='major', labelsize=6)
    ax4.set_title('Fig 4 - cost type break down', loc='left', fontsize=8, fontweight='bold')

    #scale y axis
    #both_types = rdel_bar_chart_data + cdel_bar_chart_data
    y_max = max(total_fin) + max(total_fin) * 1 / 5
    ax4.set_ylim(0, y_max)

    # size of chart and fit
    fig.canvas.draw()
    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # for title

    fig.savefig('cost_profile.png')
    plt.close()  # automatically closes figure so don't need to do manually.

    doc.add_picture('cost_profile.png', width=Inches(8))  # to place nicely in doc
    os.remove('/home/will/code/python/analysis_engine/project_summaries/cost_profile.png')

    return doc

def milestone_swimlane_charts(doc, project_name, latest_milestone_names, latest_milestone_dates, \
                              last_milestone_dates, baseline_milestone_dates, graph_title):
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    #build scatter chart
    fig, ax1 = plt.subplots()
    fig.suptitle(graph_title, fontweight='bold')  # title
    # set fig size
    # fig.set_figheight(6)
    # fig.set_figwidth(8)

    ax1.scatter(baseline_milestone_dates, latest_milestone_names, label='Baseline')
    ax1.scatter(last_milestone_dates, latest_milestone_names, label='Last Qrt')
    ax1.scatter(latest_milestone_dates, latest_milestone_names, label='Latest Qrt')

    # format the x ticks
    years = mdates.YearLocator()  # every year
    months = mdates.MonthLocator()  # every month
    years_fmt = mdates.DateFormatter('%Y')
    months_fmt = mdates.DateFormatter('%b')

    # calculate the length of the time period covered in chart. Not perfect as baseline dates can distort.
    td = (latest_milestone_dates[-1] - latest_milestone_dates[0]).days
    if td <= 365*3:
        ax1.xaxis.set_major_locator(years)
        ax1.xaxis.set_minor_locator(months)
        ax1.xaxis.set_major_formatter(years_fmt)
        ax1.xaxis.set_minor_formatter(months_fmt)
        plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45)
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight='bold')
    else:
        ax1.xaxis.set_major_locator(years)
        ax1.xaxis.set_minor_locator(months)
        ax1.xaxis.set_major_formatter(years_fmt)
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight='bold')

    ax1.legend() #insert legend

    #reverse y axis so order is earliest to oldest
    ax1 = plt.gca()
    ax1.set_ylim(ax1.get_ylim()[::-1])
    ax1.tick_params(axis='y', which='major', labelsize=7)

    #Add line of IPDC date, but only if in the time period
    if latest_milestone_dates[0] <= ipdc_date <= latest_milestone_dates[-1]:
        plt.axvline(ipdc_date)
        ax1.set_title('Line represents when IPDC will discuss Q1 20_21 portfolio management report',
                      loc='left', fontsize=8, fontweight='bold')
        #plt.text(ipdc_date, 10, 'IPDC Q1 PfM Report', rotation=90) not used as impacting tight_fit

    #size of chart and fit
    fig.canvas.draw()
    fig.tight_layout(rect=[0, 0.03, 1, 0.95]) #for title

    fig.savefig('schedule.png', bbox_inches='tight')
    plt.close() #automatically closes figure so don't need to do manually.

    doc.add_picture('schedule.png', width=Inches(8))  # to place nicely in doc
    os.remove('/home/will/code/python/analysis_engine/project_summaries/schedule.png')

    return doc

def milestone_table(doc, p_baseline_milestones, project_name):

    second_diff_data = project_time_difference(p_current_milestones, p_baseline_milestones)

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Milestone'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Change from Lst Qrt'
    hdr_cells[3].text = 'Change from BL'
    hdr_cells[4].text = 'Notes'

    # TODO specify column widths


    for milestone in p_current_milestones[project_name].keys():

        milestone_date = tuple(p_current_milestones[project_name][milestone])[0]

        try:
            if milestone_filter_start_date <= milestone_date: # filter based on date
                row_cells = table.add_row().cells
                row_cells[0].text = milestone
                if milestone_date is None:
                    row_cells[1].text = 'No date'
                else:
                    row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
                b_one_value = first_diff_data[project_name][milestone]
                row_cells[2].text = plus_minus_days(b_one_value)
                b_two_value = second_diff_data[project_name][milestone]
                row_cells[3].text = plus_minus_days(b_two_value)
                notes = p_current_milestones[project_name][milestone][milestone_date]
                row_cells[4].text = str(notes)
                paragraph = row_cells[4].paragraphs[0]
                run = paragraph.runs
                font = run[0].font
                font.size = Pt(8)  # font size = 8


        except TypeError:  # this is to deal with none types which are still placed in output
            row_cells = table.add_row().cells
            row_cells[0].text = milestone
            if milestone_date is None:
                row_cells[1].text = 'No date'
            else:
                row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
            b_one_value = first_diff_data[project_name][milestone]
            row_cells[2].text = plus_minus_days(b_one_value)
            b_two_value = second_diff_data[project_name][milestone]
            row_cells[3].text = plus_minus_days(b_two_value)
            notes = p_current_milestones[project_name][milestone][milestone_date]
            row_cells[4].text = str(notes)
            paragraph = row_cells[4].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(8)  # font size = 8

    table.style = 'Table Grid'

    make_rows_bold(table.rows[0]) # makes top of table bold. Found function on stack overflow.

    return doc

def make_rows_bold(*rows):
    '''Makes text bold in specified row'''
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def plus_minus_days(change_value):
    '''mini function to place plus or minus sign before time delta
    value in milestone_table function. Only need + signs to be added
    as negative numbers have minus already'''
    try:
        if change_value > 0:
            text = '+ ' + str(change_value)
        else:
            text = str(change_value)
    except TypeError:
        text = change_value

    return text

def get_financial_totals(project_name):
    '''gets financial data to place into the bar chart element in the financial analysis graphs'''
    key_list = [('Pre-profile RDEL',
                'Pre-profile CDEL'),
                ('Total RDEL Forecast Total',
                 'Total CDEL Forecast Total WLC'),
                ('Unprofiled RDEL Forecast Total',
                 'Unprofiled CDEL Forecast Total WLC')]

    total_cost_list = []
    rdel_cost_list = []
    cdel_cost_list = []

    index_1 = bc_index[project_name]
    index_2 = index_1[0:3]
    index_2.reverse()
    for x in index_2:
        #print(x)
        for y in key_list:
            rdel = list_of_masters_all[x].data[project_name][y[0]]
            cdel = list_of_masters_all[x].data[project_name][y[1]]
            total = rdel + cdel
            total_cost_list.append(total)
            rdel_cost_list.append(rdel)
            cdel_cost_list.append(cdel)

    return total_cost_list, rdel_cost_list, cdel_cost_list

def bar_chart_benefits(project_name):
    '''gets benefits data to place into the bar chart element in the financial analysis graphs'''

    key_list = ['Pre-profile BEN Total',
                'Unprofiled Remainder BEN Forecast - Total Monetised Benefits',
                'Total BEN Forecast - Total Monetised Benefits']

    cost_list = []

    index_1 = bc_index[project_name]
    index_2 = index_1[0:3]
    index_2.reverse()
    for x in index_2:
        #print(x)
        for y in key_list:
            ben = list_of_masters_all[x].data[project_name][y]
            #print(ben)
            cost_list.append(ben)

    return cost_list

def milestone_schedule_data(latest_m_dict, last_m_dict, baseline_m_dict, project_name):
    milestone_names = []
    mile_d_l_lst = []
    mile_d_last_lst = []
    mile_d_bl_lst = []

    #lengthy for loop designed so that all milestones and dates are stored and shown in output chart, even if they
    #were not present in last and baseline data reporting
    for m in list(latest_m_dict[project_name].keys()):
        if m == 'Project - Business Case End Date':  # filter out as to far in future
            pass
        else:
            if m is not None:
                m_d = tuple(latest_m_dict[project_name][m])[0]

            if m in list(last_m_dict[project_name].keys()):
                m_d_lst = tuple(last_m_dict[project_name][m])[0]
            else:
                m_d_lst = tuple(latest_m_dict[project_name][m])[0]

            if m in list(baseline_m_dict[project_name].keys()):
                m_d_bl = tuple(baseline_m_dict[project_name][m])[0]
            else:
                m_d_bl = tuple(latest_m_dict[project_name][m])[0]

            if m_d is not None:
                milestone_names.append(m)
                mile_d_l_lst.append(m_d)
                if m_d_lst is not None:
                    mile_d_last_lst.append(m_d_lst)
                else:
                    mile_d_last_lst.append(m_d)
                if m_d_bl is not None:
                    mile_d_bl_lst.append(m_d_bl)
                else:
                    if m_d_lst is not None:
                        mile_d_bl_lst.append(m_d_lst)
                    else:
                        mile_d_bl_lst.append(m_d)

    return milestone_names, mile_d_l_lst, mile_d_last_lst, mile_d_bl_lst

def milestone_schedule_data_filtered(latest_m_dict, last_m_dict, baseline_m_dict, project_name):
    milestone_names = []
    mile_d_l_lst = []
    mile_d_last_lst = []
    mile_d_bl_lst = []

    #lengthy for loop designed so that all milestones and dates are stored and shown in output chart, even if they
    #were not present in last and baseline data reporting
    for m in list(latest_m_dict[project_name].keys()):
        if m == 'Project - Business Case End Date':  # filter out as to far in future
            pass
        else:
            if m is not None:
                m_d = tuple(latest_m_dict[project_name][m])[0]

            if m in list(last_m_dict[project_name].keys()):
                m_d_lst = tuple(last_m_dict[project_name][m])[0]
            else:
                m_d_lst = tuple(latest_m_dict[project_name][m])[0]

            if m in list(baseline_m_dict[project_name].keys()):
                m_d_bl = tuple(baseline_m_dict[project_name][m])[0]
            else:
                m_d_bl = tuple(latest_m_dict[project_name][m])[0]

            if m_d is not None:
                if milestone_filter_start_date <= m_d <= milestone_filter_end_date:
                    milestone_names.append(m)
                    mile_d_l_lst.append(m_d)
                    if m_d_lst is not None:
                        mile_d_last_lst.append(m_d_lst)
                    else:
                        mile_d_last_lst.append(m_d)
                    if m_d_bl is not None:
                        mile_d_bl_lst.append(m_d_bl)
                    else:
                        if m_d_lst is not None:
                            mile_d_bl_lst.append(m_d_lst)
                        else:
                            mile_d_bl_lst.append(m_d)

    return milestone_names, mile_d_l_lst, mile_d_last_lst, mile_d_bl_lst


'''RUNNING PROGRAMME'''

'''enter into the printing function the quarter details for the output files e.g. _q4_1920 (note put underscore at 
front)'''
produce_word_doc()


